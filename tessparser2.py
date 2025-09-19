import fitz
from unidecode import unidecode
from docx import Document
from PIL import Image
import pytesseract
import json
import os
import re
from difflib import SequenceMatcher

# Ensure tesseract is installed and path is set if necessary
pytesseract.pytesseract.tesseract_cmd = r'C:\Users\GAURAV D\AppData\Local\Programs\Tesseract-OCR\tesseract.exe'

# Input file (assuming the same path)
file_path = r'C:\Users\GAURAV D\Desktop\ai_startup_hackathon\Company Data\14. Ziniosa\Ziniosa Pitch Deck.pdf'

# Output folder
output_folder = r'C:\Users\GAURAV D\Desktop\ai_startup_hackathon\processed_output'
os.makedirs(output_folder, exist_ok=True)

def extract_text_in_visual_order(page, y_tolerance=5):
    """
    Extracts text from a PDF page, attempting to maintain visual reading order.
    Groups words into lines based on y-coordinate tolerance, then sorts lines.
    """
    words = page.get_text("words")  # (x0, y0, x1, y1, word, block_no, line_no, word_no)
    if not words:
        return ""

    # Sort words primarily by y0, then by x0
    # Store them as tuples (x0, y0, x1, y1, text) for easier manipulation
    sorted_words = [(w[0], w[1], w[2], w[3], unidecode(w[4]).replace('\n', ' ').strip()) for w in words]
    sorted_words = [w for w in sorted_words if w[4]] # Filter out empty strings
    sorted_words.sort(key=lambda w: (w[1], w[0]))

    lines = []
    current_line_words = []

    for i, word_info in enumerate(sorted_words):
        x0, y0, x1, y1, text = word_info

        if not current_line_words:
            current_line_words.append(word_info)
        else:
            # Check if the current word is on the same line as the first word in current_line_words
            first_word_y0_in_current_line = current_line_words[0][1]
            if abs(y0 - first_word_y0_in_current_line) <= y_tolerance:
                current_line_words.append(word_info)
            else:
                # New line detected, process the current_line_words
                current_line_words.sort(key=lambda w: w[0]) # Sort by x0 within the line
                lines.append(" ".join([w[4] for w in current_line_words]))
                current_line_words = [word_info] # Start a new line with the current word

    # Add the last accumulated line
    if current_line_words:
        current_line_words.sort(key=lambda w: w[0]) # Sort by x0 within the line
        lines.append(" ".join([w[4] for w in current_line_words]))

    return "\n".join(lines)

def clean_ocr_text(text):
    """
    Advanced cleaning for OCR text:
    - Removes common OCR artifacts (e.g., deg, single chars, repeated chars)
    - Removes excessive newlines and spaces
    - Filters out very short, non-alphanumeric lines
    """
    text = unidecode(text)
    
    # Replace common OCR noise patterns and special characters that don't add meaning
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text) # Remove control characters
    text = re.sub(r'[\u201c\u201d\u2018\u2019]', "'", text) # Normalize quotes
    text = re.sub(r'[\u2013\u2014]', '-', text) # Normalize dashes
    text = re.sub(r'[\u2022\u25cf\u25cb]', '*', text) # Normalize bullet points

    # Remove repeated characters often seen in OCR (e.g., '==', '---', ';;;')
    text = re.sub(r'(.)\1{3,}', r'\1', text) # Any char repeated 3+ times
    
    # Remove strings like "deg", "O", "Ul", "i", "L", "SS", "PS", "SB", "Ee", "aa", "==--" if they appear as isolated tokens
    # Using word boundaries for safety. 'deg' could be part of a word.
    text = re.sub(r'\b(deg|O|Ul|i|L|SS|PS|SB|Ee|aa)\b', '', text, flags=re.IGNORECASE)
    text = re.sub(r'--[->]', '', text) # Remove specific arrow-like noise
    text = re.sub(r'[ioIO]\}', '', text) # Remove common bracket-like noise from bullet points
    text = re.sub(r'[\"\'\`]', '', text) # Remove common quote artifacts
    text = re.sub(r'\.{2,}', '.', text) # Replace multiple dots with a single one

    # Remove lines that are mostly noise (e.g., just punctuation or very few characters)
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        stripped_line = line.strip()
        # Keep lines if they have at least one alphanumeric character OR are longer than a minimal length
        if re.search(r'[a-zA-Z0-9]', stripped_line) or len(stripped_line) > 5:
            cleaned_lines.append(stripped_line)
    text = '\n'.join(cleaned_lines)
    
    # Remove excessive newlines and spaces, replace multiple spaces with single space
    text = re.sub(r'\n\s*\n', '\n\n', text)
    text = re.sub(r' +', ' ', text)
    text = text.strip()
    
    return text

def is_similar(s1, s2, threshold=0.9):
    """Checks if two strings are similar above a given threshold."""
    return SequenceMatcher(None, s1, s2).ratio() > threshold

def extract_full_text_from_document(file_path, use_ocr_for_all_pages=True):
    pages_text = {}
    total_word_count = 0

    try:
        if file_path.endswith('.pdf'):
            doc = fitz.open(file_path)
            for page_num, page in enumerate(doc, start=1):
                # 1. Native text extraction in visual order
                native_text = extract_text_in_visual_order(page)

                # 2. OCR extraction for potentially missed content (like logos, text in images)
                ocr_text = ""
                if use_ocr_for_all_pages:
                    pix = page.get_pixmap(dpi=300) # Higher DPI for better OCR
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    ocr_raw_text = pytesseract.image_to_string(img)
                    ocr_text = clean_ocr_text(ocr_raw_text)

                final_page_text_parts = []
                native_lines = [line.strip() for line in native_text.split('\n') if line.strip()]
                ocr_lines = [line.strip() for line in ocr_text.split('\n') if line.strip()]
                
                # Add native text first, as it's generally more accurate
                if native_lines:
                    final_page_text_parts.extend(native_lines)

                # Convert existing text to a lowercased, space-tokenized set for quick lookup
                existing_words_lower = set(re.findall(r'\b\w+\b', ("\n".join(final_page_text_parts)).lower()))
                
                # Track if 'naario' has been added to this page's text
                naario_added_for_page = False
                if "naario" in existing_words_lower:
                    naario_added_for_page = True

                # Iterate through OCR lines to add non-duplicate content
                for ocr_line in ocr_lines:
                    is_duplicate = False
                    # Check for exact or highly similar line duplication with native text
                    for native_line in native_lines:
                        if is_similar(native_line.lower(), ocr_line.lower(), threshold=0.9): # Stricter threshold for full lines
                            is_duplicate = True
                            break
                    
                    if not is_duplicate:
                        # Further check for common, short elements like the 'naario' logo
                        if "naario" in ocr_line.lower():
                            if not naario_added_for_page:
                                final_page_text_parts.append("naario") # Standardize to lowercase if it's the logo
                                naario_added_for_page = True
                            # If 'naario' is in the line, and we've already added it, don't add the full line unless it's very distinct
                            if is_similar(ocr_line.lower(), "naario", threshold=0.95): # If the whole line IS "naario", we've handled it
                                continue
                            
                        # For other OCR content, add if it contains significant new words
                        ocr_line_words_lower = set(re.findall(r'\b\w+\b', ocr_line.lower()))
                        new_words = ocr_line_words_lower - existing_words_lower
                        
                        # Add the OCR line if it contains a reasonable number of new words, or it's a longer line
                        # This heuristic tries to capture new info (like product names) without blindly appending.
                        if len(new_words) > 2 or (len(ocr_line) > 30 and len(new_words) > 0): # At least 3 new words, or a long line with some new words
                            final_page_text_parts.append(ocr_line)
                            # Update existing_words_lower for subsequent checks in the same page
                            existing_words_lower.update(ocr_line_words_lower)
                
                # Final cleaning and re-joining
                final_page_text = "\n".join(final_page_text_parts)
                final_page_text = clean_ocr_text(final_page_text) # Final clean after combination

                if final_page_text.strip():
                    cleaned_text = final_page_text # Already unidecoded and cleaned
                    word_count = len(cleaned_text.split())
                    total_word_count += word_count
                    pages_text[f"page_{page_num}"] = {
                        "word_count": word_count,
                        "text": cleaned_text
                    }

            doc.close()

        elif file_path.endswith('.docx'):
            document = Document(file_path)
            doc_text = []
            for paragraph in document.paragraphs:
                text = unidecode(paragraph.text).strip()
                if text:
                    doc_text.append(text)

            for table in document.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        text = unidecode(cell.text).strip()
                        if text:
                            row_text.append(text)
                    if row_text:
                        doc_text.append(' | '.join(row_text))

            if doc_text:
                full_text = '\n\n'.join(doc_text)
                word_count = len(full_text.split())
                total_word_count = word_count
                pages_text["page_1"] = {
                    "word_count": word_count,
                    "text": full_text
                }

        return pages_text, total_word_count if pages_text else (None, 0)

    except Exception as e:
        print(f"Error extracting text from file: {file_path}, Error: {e}")
        return None, 0


if __name__ == "__main__":
    pages_dict, total_word_count = extract_full_text_from_document(file_path, use_ocr_for_all_pages=True)
    if pages_dict:
        output_data = {
            "document_name": os.path.basename(file_path),
            "total_word_count": total_word_count,
            "pages": pages_dict
        }

        base_name = os.path.splitext(os.path.basename(file_path))[0]
        output_file = os.path.join(output_folder, f"{base_name}_pages.json")

        with open(output_file, "w", encoding="utf-8") as f:
            json.dump(output_data, f, indent=4, ensure_ascii=False)

        print(f" Page-wise text with metadata saved to {output_file}")
    else:
        print("No text extracted.")