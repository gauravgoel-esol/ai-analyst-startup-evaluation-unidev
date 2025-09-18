import os
import fitz 
import json
import csv
from unidecode import unidecode
from docx import Document 
import google.generativeai as genai
from dotenv import load_dotenv 
import time # For retry mechanism

# Load environment variables from .env file
load_dotenv()

# Get the API key from environment variable
gemini_api_key = os.getenv('GEMINI_API_KEY')

# Function to extract full text from a document (PDF or DOCX)
def extract_full_text_from_document(file_path):
    full_text = []
    try:
        if file_path.endswith('.pdf'):
            doc = fitz.open(file_path)
            for page in doc:
                # Extract text block by block, joining lines and spans
                # Less aggressive cleaning than resume parser, keeping newlines as spaces
                for block in page.get_text('dict')['blocks']:
                    if block['type'] == 0:  # Text block
                        block_text = []
                        for line in block['lines']:
                            line_text = []
                            for span in line['spans']:
                                # Convert to ASCII, replace internal newlines, but keep spaces
                                text = unidecode(span['text']).replace('\n', ' ').strip()
                                if text:
                                    line_text.append(text)
                            if line_text:
                                block_text.append(' '.join(line_text))
                        if block_text:
                            full_text.append('\n'.join(block_text)) # Use newline to separate blocks/paragraphs
            doc.close()

        elif file_path.endswith('.docx'):
            document = Document(file_path)
            # Extract text from paragraphs
            for paragraph in document.paragraphs:
                text = unidecode(paragraph.text).strip()
                if text:
                    full_text.append(text)
            
            # Extract text from tables (if any)
            for table in document.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        text = unidecode(cell.text).strip()
                        if text:
                            row_text.append(text)
                    if row_text:
                        full_text.append(' | '.join(row_text)) # Use | to separate cell content in a row

        return '\n\n'.join(full_text) if full_text else None
    except Exception as e:
        print(f"Error extracting text from file: {file_path}, Error: {e}")
        return None

# Function to generate a comprehensive JSON analysis from document text
def generate_document_analysis_json(document_text, file_name, max_retries=3, initial_delay=5):
    # This prompt is designed to extract a wide range of information and structure it
    # in a flexible JSON format.
    format_prompt = f'''
    You are an expert document analyst capable of extracting structured information from various document types, particularly business plans and presentations.
    Review the provided document text for "{file_name}" and extract all relevant details, presenting them in a comprehensive JSON format.
    Ensure all numerical values are extracted as numbers (integers or floats) where appropriate (e.g., "1.5 Cr" should be 15000000.0, "64%" should be 0.64).
    If specific fields are not found or applicable, they should be omitted or kept as empty strings/arrays, or null for single values.

    Based on the document data, provide the following details in JSON format:

    {{
        "document_name": "{file_name}",
        "document_type": "[Identify the specific type of document, e.g., 'Business Plan', 'Investor Deck', 'Company Report', 'Market Analysis', 'Founders Checklist', etc. Choose the most precise fit.]",
        "title": "[Main title of the document. If it's an investor deck, infer a title like 'Naario Investor Deck' or 'Naario Business Presentation' if not explicitly stated on the first page. Otherwise, extract the most prominent title.]",
        "summary": "[A concise summary of the document's main objective, key business insights, and overall conclusions. Focus on Naario's mission, strategy, and market position.]",
        "key_themes": [], // List of the main topics or strategic pillars discussed, e.g., "Millet-first ecosystem", "Community-led growth", "Sustainability", "Financial Projections"
        "identified_entities": {{
            "organizations": [], // List all company names, brands, platforms, institutions mentioned, e.g., "Bigbasket", "Amazon", "Rupyz", "HORECA", "TSUV", "Meta", "Blue Tokai", "Tata Soulfull"
            "people": [],      // List all named individuals, e.g., "Anamika Pandey", "Charul Chandak". If a founder profile is described but the name is not in the text, leave name null.
            "locations": [],   // List all cities, states, countries, specific addresses mentioned, e.g., "India", "Bengaluru", "Gorakhpur", "NIT Warangal", "Chhatarpur"
            "dates": [],       // List all significant years, specific dates, or fiscal year mentions, e.g., "2019", "FY 2025-26", "2023"
            "products_services": [] // List specific product names or services mentioned, e.g., "Millet Attas", "Muesli", "Snack Bhakhri", "Rupyz SaaS platform", "Shopify"
        }},
        "sections": [ // Extract main sections with their headings, a brief content summary, and key points/bullet points.
            {{
                "heading": "[Main heading of the section]",
                "content_summary": "[A brief summary of this section's content]",
                "key_points": [] // List of important points, sub-headings, or key takeaways within the section
            }}
            // ... more sections as found in the document
        ],
        "tables_data": [ // IMPORTANT: Identify and extract data from any tabular structures within the document text. This includes explicitly formatted tables or data presented in a clear row/column format (like competitive analysis, financial breakdowns, etc.).
            {{
                "title": "[Title or context of the table, e.g., 'Competitor Analysis Framework']",
                "headers": [], // List of column headers (e.g., "Category", "Company Name", "Headquarters")
                "rows": [[]]   // List of rows, where each row is a list of cell values in the order of headers
            }}
            // ... more tables if found
        ],
        "current_financials": {{ // Current financial metrics as of the document's context
            "monthly_recurring_revenue": null, // Convert "INR 14 Lakhs" to 1400000.0
            "annual_recurring_revenue": null, // Convert "INR 1.5 Cr" to 15000000.0
            "monthly_burn": null, // Convert "INR 15k" to 15000.0
            "runway_months": null, // Convert "8 Months" to 8 (integer)
            "gross_margin_avg": null, // Convert "64%" to 0.64
            "gross_margin_range": "", // e.g., "58-76%"
            "contribution_margin_1_percent": null, // Convert "64%" to 0.64
            "contribution_margin_2_percent": null, // Convert "29%" to 0.29
            "contribution_margin_3_percent": null  // Convert "18%" to 0.18
        }},
        "financial_projections": {{ // If the document contains future financial projections by fiscal year (e.g., from tables)
            "revenue_projections": {{ // Example for FYs, adjust if different fiscal years are present
                "FY_2025-26": null,
                "FY_2026-27": null,
                "FY_2027-28": null
            }},
            "gross_margin_projections": {{
                "FY_2025-26": null,
                "FY_2026-27": null,
                "FY_2027-28": null
            }},
            "ebitda_projections": {{
                "FY_2025-26": null,
                "FY_2026-27": null,
                "FY_2027-28": null
            }}
        }},
        "funding_details": {{ // Details related to fundraising efforts
            "total_funding_raised": null, // Convert "1.85 Cr" to 18500000.0
            "funding_ask_min": null, // Convert "50 Lakhs" to 5000000.0
            "funding_ask_max": null, // Convert "1 Cr" to 10000000.0
            "structure": "", // e.g., "CCPS", "SAFE", "Convertible Note"
            "valuation_pre_money": null, // Convert "23 Cr" to 23000000.0 (if explicit "Pre-Money" value)
            "valuation_cap": null, // Convert "29 Cr" to 29000000.0
            "valuation_floor": null, // Convert "23 Cr" to 23000000.0
            "lead_investor": "",
            "incoming_investors": [], // List of investor names
            "existing_investors": [], // List of investor names
            "current_commitments_percent": null // Convert "40%" to 0.40
        }},
        "funding_allocation": {{ // If funding allocation breakdown is provided as percentages, convert to floats.
            "scaling_distribution_percent": null, // e.g., 0.13
            "brand_awareness_customer_acquisition_percent": null, // e.g., 0.30
            "supply_chain_production_scale_up_percent": null,
            "hiring_team_expansion_percent": null,
            "product_innovation_r_d_percent": null,
            "technology_data_driven_expansion_percent": null
        }},
        "team_members": [ // Information about key team members, especially founders. Infer roles if not explicit.
            {{
                "name": "[Full name of the team member if explicitly mentioned, e.g., 'Anamika Pandey', 'Charul Chandak'. If not named in the document text, leave as null.]",
                "role": "", // e.g., "Founder", "Head of New Initiatives", "Partners Lead"
                "previous_experience": "", // Summarize previous work history from the text
                "education": "" // Summarize educational background from the text
            }}
        ],
        "key_business_narratives": {{ // Sections describing the core business aspects in detail
            "problem_solved": "[Summarize the key problem Naario aims to solve for its target market, as described in the text.]",
            "business_model_description": "[Explain Naario's core business model, how it operates, and generates value, as described in the text.]",
            "sales_pipeline": "[Details about the current sales pipeline value and growth targets mentioned in the text.]",
            "projected_growth_opportunities": [], // List specific areas for future growth (e.g., "Quick Commerce", "B2B Pipeline spanning across Chaayos, PVR, Cinepolis")
            "market_trends": "[Summarize the relevant market trends driving Naario's opportunity, as described in the text.]",
            "competitive_edge": "[Describe Naario's unique advantages and differentiators, as outlined in the text.]",
            "urgency_opportunity": "[Explain why now is the critical time for Naario's growth and market capture, as presented in the text.]"
        }},
        "risks_and_mitigation": [ // List identified risks and their corresponding mitigation strategies as pairs.
            {{
                "risk": "",
                "mitigation_strategy": ""
            }}
        ],
        "facilities": {{ // Details about physical infrastructure
            "office_details": "",
            "plant_details": "",
            "warehouses": ""
        }},
        "technology_stack": {{ // Details about technology used
            "b2b2c_distribution_platform": "",
            "d2c_front_end_platform": "",
            "customer_connect_support_platform": ""
        }}
    }}
    '''
   
    prompt = f'''{format_prompt}

    Here is the document text for "{file_name}" that you need to analyze:
    "{document_text}"
    '''
   
    for attempt in range(max_retries):
        try:
            model = genai.GenerativeModel('gemini-1.5-flash', generation_config={"response_mime_type": "application/json", "temperature": 0.3}) # Added temperature
            response = model.generate_content(prompt, request_options={"timeout": 600}) # Increased timeout
            result = response.text
            # Attempt to parse immediately to catch JSON errors before returning
            json.loads(result) 
            return result
        except json.JSONDecodeError as e:
            print(f"Attempt {attempt + 1}/{max_retries}: JSON decoding failed for {file_name}. Error: {e}. Retrying...")
            print(f"Raw response (first 1000 chars): {result[:1000]}...") # Print more chars for debugging
            time.sleep(initial_delay * (attempt + 1)) # Exponential backoff
        except Exception as e:
            print(f"Attempt {attempt + 1}/{max_retries}: Error generating JSON for {file_name}: {e}. Retrying...")
            if hasattr(e, 'response'):
                print(f"API Response Status: {e.response.status_code}")
                print(f"API Response Text: {e.response.text}")
            time.sleep(initial_delay * (attempt + 1)) # Exponential backoff
    
    print(f"Failed to generate valid JSON after {max_retries} attempts for {file_name}.")
    return None

# --- Main execution script ---

# Adjust this path as needed. If you want to process all subfolders recursively,
# you would need to use os.walk instead of os.listdir.
# For the Naario example, let's assume 'Company Data/01. Data stride/02. Naario' is the target folder.
folder_path = r'C:\Users\GAURAV D\Desktop\ai_startup_hackathon\Company Data\02. Naario' # Corrected path

# Create an output folder for JSON results
output_json_folder = 'processed_documents_json'
os.makedirs(output_json_folder, exist_ok=True)

# Configure the Generative AI API and check if the key is present
if not gemini_api_key:
    print("Error: Gemini API key not found. Please set GEMINI_API_KEY in your .env file.")
    exit(1) # Exit with an error code
else:
    genai.configure(api_key=gemini_api_key)
    print("Gemini API key successfully loaded and configured.")
    
    # Process each document in the folder
    for file_name in os.listdir(folder_path):
        # Skip temporary Word files and any other hidden/system files
        if file_name.startswith('~') or file_name.startswith('.') or file_name.lower().endswith('.tmp'): 
            print(f"Skipping temporary or hidden file: {file_name}")
            continue

        file_path = os.path.join(folder_path, file_name)
        
        if file_name.endswith(('.pdf', '.docx')):
            print(f"\nProcessing document: {file_name}")
            document_text = extract_full_text_from_document(file_path)
            
            if document_text:
                document_json_str = generate_document_analysis_json(document_text, file_name)
                if document_json_str:
                    try:
                        # Attempt to parse and pretty-print the JSON
                        document_dict = json.loads(document_json_str)
                        output_file_name = os.path.splitext(file_name)[0] + '.json'
                        output_file_path = os.path.join(output_json_folder, output_file_name)
                        
                        with open(output_file_path, mode='w', encoding='utf-8') as json_file:
                            json.dump(document_dict, json_file, indent=4, ensure_ascii=False)
                        print(f"Successfully processed and saved JSON for: {file_name} to {output_file_path}")
                    except json.JSONDecodeError as e:
                        print(f"Fatal Error: JSON decoding failed after retries for {file_name}: {e}. Raw response (first 1000 chars): {document_json_str[:1000]}...")
                    except Exception as e:
                        print(f"An unexpected error occurred while saving JSON for {file_name}: {e}")
                else:
                    print(f"Error: JSON analysis could not be generated after retries for document: {file_name}")
            else:
                print(f"Could not extract text from document: {file_name}")
        else:
            print(f"Skipping unsupported file format: {file_name}")

    print("\nDocument processing complete!")